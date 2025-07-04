from docx import Document
import os
import uuid

def generate_project_doc(summary: str) -> str:
    document = Document()

    document.add_heading("Project Proposal", 0)

    document.add_heading("Title:", level=1)
    document.add_paragraph("AI-Powered Meeting Summarizer and Project Planner")

    document.add_heading("Abstract:", level=1)
    document.add_paragraph(
        "This project aims to transform meeting transcripts into actionable, structured project documentation using AI-based summarization and document generation."
    )

    document.add_heading("Summary:", level=1)
    document.add_paragraph(summary)

    document.add_heading("Technologies Likely to Be Used:", level=1)
    document.add_paragraph("- Python\n- FastAPI\n- Hugging Face Transformers / OpenAI API\n"
                           "- Google Docs API or DOCX Export\n- Chrome Extension APIs (for Google Meet integration)\n"
                           "- Optional: OCR, Speech-to-Text, Task Scheduler")

    document.add_heading("Project Tasks:", level=1)
    document.add_paragraph(
        "1. Audio recording from meetings\n"
        "2. Transcription using Whisper or alternative\n"
        "3. Summarization using BART or GPT-4\n"
        "4. Document generation (DOCX / Google Docs)\n"
        "5. Chrome extension UI\n"
        "6. Export features\n"
        "7. Future: Master Control Agent"
    )

    document.add_heading("Deliverables:", level=1)
    document.add_paragraph(
        "- Meeting transcript\n- Summary\n- Project document\n- API endpoints\n- Chrome extension interface"
    )

    document.add_paragraph("\nGenerated by AI based on summarized meeting content.")

    # Save the file
    output_dir = "docs"
    os.makedirs(output_dir, exist_ok=True)
    filename = f"{output_dir}/project_doc_{uuid.uuid4().hex}.docx"
    document.save(filename)

    return filename
